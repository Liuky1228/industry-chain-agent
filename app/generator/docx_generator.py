"""Word 报告生成模块（P1 重构 — 9章标准结构）

使用 python-docx 生成机构视角的产业链分析报告。
章节结构遵循产业研究标准：

  第一章：产业链全景概览（定义、流转、市场规模）= 原1 + 原5
  第二章：产业链深度拆解（上/中/下游分节）= 原3
  第三章：产业逻辑分析 = 原2 + 原4
  第四章：产业未来趋势与风险 = 原6 + 原8
  第五章：近期产业链动态 = 原9
  附录：重点企业分析、企业关系
  免责声明
"""

import os
import re
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def generate_report(chain_data: dict, output_dir: str = "data/reports") -> str:
    """
    生成产业链分析报告 Word 文档（标准结构：5 章正文 + 附录）

    Args:
        chain_data: 产业链分析完整结果 (merge_and_analyze 的返回值)
        output_dir: 输出目录

    Returns:
        生成的文件路径
    """
    os.makedirs(output_dir, exist_ok=True)

    industry_name = chain_data.get("industry_name", "未知产业")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{industry_name}_产业链分析报告_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)

    doc = Document()

    # ── 全局样式 ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # ── 封面 ──
    _add_cover(doc, industry_name)

    # ── 目录 ──
    doc.add_page_break()
    _add_toc(doc)

    # ── 获取叙述数据（如有）──
    narratives = chain_data.get("_narratives", {})

    # ── 第一章：产业链全景概览（定义、流转、市场规模）= 原1 + 原5 ──
    _add_chapter_boundary(doc, chain_data, narratives)
    _add_chapter_profile(doc, chain_data, narratives)

    # ── 第二章：产业链深度拆解 = 原3 ──
    _add_chapter_node_analysis(doc, chain_data, narratives)

    # ── 第三章：产业逻辑分析 = 原2 + 原4 ──
    _add_chapter_causal_logic(doc, chain_data, narratives)
    _add_chapter_transmission(doc, chain_data, narratives)

    # ── 第四章：产业未来趋势与风险 = 原6 + 原8 ──
    _add_chapter_trends(doc, chain_data, narratives)
    _add_chapter_risk(doc, chain_data, narratives)

    # ── 第五章：近期产业链动态 = 原9 ──
    _add_chapter_events(doc, chain_data)

    # ── 附录：重点企业分析 + 企业关系（与前端 / HTML 一致）──
    _add_key_companies(doc, chain_data)
    _add_relations(doc, chain_data)

    # ── 免责声明 ──
    _add_disclaimer(doc)

    # ── 保存 ──
    doc.save(filepath)
    logger.info(f"报告已生成（5章标准结构）: {filepath}")
    return filepath


# ══════════════════════════════════════════════════════════════
# ── 封面 & 目录 ──
# ══════════════════════════════════════════════════════════════

def _add_cover(doc: Document, industry_name: str):
    """添加封面"""
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{industry_name}\n产业链分析报告")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"AI 智能分析 | {datetime.now().strftime('%Y年%m月%d日')}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 副标题：机构视角
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("产业研究机构 · 内部参考")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True


def _add_toc(doc: Document):
    """添加目录（5章标准结构）"""
    doc.add_heading("目录", level=1)
    chapters = [
        "第一章  产业链全景概览（定义、流转、市场规模）",
        "第二章  产业链深度拆解",
        "    2.1 上游节点分析",
        "    2.2 中游节点分析",
        "    2.3 下游节点分析",
        "第三章  产业逻辑分析",
        "第四章  产业未来趋势与风险",
        "第五章  近期产业链动态",
        "附录    重点企业分析",
        "附录    企业关系",
    ]
    for ch in chapters:
        p = doc.add_paragraph(ch)
        p.paragraph_format.space_after = Pt(6)


# ══════════════════════════════════════════════════════════════
# ── 第一章：产业链全景概览（定义、流转、市场规模）= 原1 + 原5 ──
# ══════════════════════════════════════════════════════════════

def _add_chapter_boundary(doc: Document, data: dict, narratives: dict):
    """第一章：产业链全景概览（定义、流转、市场规模）= 原1 + 原5"""
    doc.add_heading("第一章  产业链全景概览（定义、流转、市场规模）", level=1)

    # 产业定义与覆盖范围
    desc = data.get("industry_description", "")
    if desc:
        doc.add_heading("产业定义与覆盖范围", level=2)
        doc.add_paragraph(desc)

    # AI 叙述（如有）
    narrative = narratives.get("industry_boundary", "")
    if narrative:
        doc.add_paragraph(narrative)

    # 核心主干环节列表
    segments = data.get("chain_segments", {})
    core_segments = []
    for level in ["upstream", "midstream", "downstream"]:
        for seg in segments.get(level, []):
            core_segments.append((seg.get("segment_name", ""), level))

    if core_segments:
        doc.add_heading("核心主干环节", level=2)
        level_map = {"upstream": "上游", "midstream": "中游", "downstream": "下游"}

        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = "序号"
        hdr[1].text = "环节名称"
        hdr[2].text = "产业链位置"

        for idx, (seg_name, level) in enumerate(core_segments, 1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = seg_name
            row[2].text = level_map.get(level, level)


# ══════════════════════════════════════════════════════════════
# ── 第三章：产业逻辑分析 = 原2 + 原4 ──
# ══════════════════════════════════════════════════════════════

def _add_chapter_causal_logic(doc: Document, data: dict, narratives: dict):
    """第三章：产业逻辑分析 = 原2 + 原4"""
    doc.add_heading("第三章  产业逻辑分析", level=1)

    # P0 字段：chain_causal_logic
    causal_logic = data.get("chain_causal_logic", "")
    if causal_logic:
        doc.add_heading("产业链构成逻辑", level=2)
        doc.add_paragraph(causal_logic)

    # AI 叙述（如有）
    narrative = narratives.get("causal_logic", "")
    if narrative:
        doc.add_paragraph(narrative)

    # 产业链流转路径
    flow = data.get("chain_flow", "")
    if flow:
        doc.add_heading("产业链流转路径", level=2)
        p = doc.add_paragraph()
        run = p.add_run(flow)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

    # 如果没有因果逻辑也没有叙述，给个提示
    if not causal_logic and not narrative:
        doc.add_paragraph(
            "（注：产业链成因逻辑尚未生成。可通过重新执行分析任务，"
            "系统将自动生成产业链结构成因的深度分析。）"
        )


# ══════════════════════════════════════════════════════════════
# ── 第二章：产业链深度拆解 = 原3 ──
# ══════════════════════════════════════════════════════════════

def _add_chapter_node_analysis(doc: Document, data: dict, narratives: dict):
    """第二章：产业链深度拆解（上/中/下游分节）= 原3"""
    doc.add_heading("第二章  产业链深度拆解", level=1)

    level_configs = [
        ("upstream", "3.1 上游节点分析", "node_analysis_upstream"),
        ("midstream", "3.2 中游节点分析", "node_analysis_midstream"),
        ("downstream", "3.3 下游节点分析", "node_analysis_downstream"),
    ]

    for level_key, section_title, narrative_key in level_configs:
        doc.add_heading(section_title, level=2)

        # AI 叙述（如有）
        narrative = narratives.get(narrative_key, "")
        if narrative:
            doc.add_paragraph(narrative)
            doc.add_paragraph()  # 间距

        # 环节详细分析
        segments = data.get("chain_segments", {}).get(level_key, [])
        if not segments:
            doc.add_paragraph("暂无该层级的详细数据。")
            continue

        for seg in segments:
            seg_name = seg.get("segment_name", "未命名环节")
            doc.add_heading(seg_name, level=3)

            # 环节描述
            desc = seg.get("description", "")
            if desc:
                doc.add_paragraph(desc)

            # P0: 四维度属性表
            _add_segment_attributes_table(doc, seg)

            # 集中度
            concentration = seg.get("concentration", "")
            if concentration:
                doc.add_paragraph(f"市场集中度：{concentration}")

            # 代表企业表格
            _add_segment_companies_table(doc, data, seg)

            doc.add_paragraph()  # 段间距


def _add_segment_attributes_table(doc: Document, seg: dict):
    """添加环节四维度属性表（P0 新增）"""
    attrs = {
        "功能定位": seg.get("functional_position", ""),
        "价值占比": seg.get("value_proportion", ""),
        "技术壁垒": seg.get("technical_barrier", ""),
        "竞争格局": seg.get("competitive_landscape", ""),
    }

    # 只有至少一个维度有值时才生成表格
    if not any(attrs.values()):
        return

    doc.add_paragraph("环节核心属性：")
    table = doc.add_table(rows=len(attrs), cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(attrs.items()):
        row = table.rows[i].cells
        row[0].text = label
        row[1].text = value or "（待补充）"
        # 标签列加粗
        for paragraph in row[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        # 值列字号
        for paragraph in row[1].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)


def _add_segment_companies_table(doc: Document, data: dict, seg: dict):
    """添加环节代表企业表格（仅列出股票代码与主营业务均非空的企业）"""
    companies_in_seg = seg.get("companies", [])
    if not companies_in_seg:
        return

    companies_map = {
        c["name"]: c for c in data.get("companies", [])
        if c.get("name") and isinstance(c["name"], str) and c["name"].strip()
    }

    # 先筛选：股票代码与主营业务均非空才列出
    valid_rows = []
    for comp_name in companies_in_seg:
        if not comp_name or not isinstance(comp_name, str):
            continue
        comp_info = companies_map.get(comp_name.strip(), {})
        stock_code = (comp_info.get("stock_code") or "").strip()
        main_business = (comp_info.get("main_business") or "").strip()
        if not stock_code or not main_business:
            continue
        valid_rows.append((comp_name, stock_code, main_business))

    if not valid_rows:
        return

    doc.add_paragraph("代表企业：")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    hdr[0].text = "企业名称"
    hdr[1].text = "股票代码"
    hdr[2].text = "主营业务"

    for comp_name, stock_code, main_business in valid_rows:
        row = table.add_row().cells
        row[0].text = comp_name
        row[1].text = stock_code
        row[2].text = main_business[:50]


# ══════════════════════════════════════════════════════════════
# ── 第三章（下）：价值传导路径分析 ──
# ══════════════════════════════════════════════════════════════

def _add_chapter_transmission(doc: Document, data: dict, narratives: dict):
    """第三章（续）：产业价值传导路径分析"""
    # 合并入第三章：产业逻辑分析（不再单独设章标题）

    # AI 叙述（如有）
    narrative = narratives.get("transmission_path", "")
    if narrative:
        doc.add_paragraph(narrative)
        doc.add_paragraph()

    # P0: 传导关系表
    transmissions = data.get("transmission_relations", [])
    if transmissions:
        doc.add_heading("环节间传导关系", level=2)

        # 按传导类型分组
        type_groups = {}
        for tr in transmissions:
            t_type = tr.get("transmission_type", "其他")
            type_groups.setdefault(t_type, []).append(tr)

        for t_type, items in type_groups.items():
            doc.add_heading(f"{t_type}路径", level=3)

            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            hdr[0].text = "起始环节"
            hdr[1].text = "目标环节"
            hdr[2].text = "传导逻辑描述"

            for item in items:
                row = table.add_row().cells
                row[0].text = item.get("from_segment", "")
                row[1].text = item.get("to_segment", "")
                row[2].text = item.get("description", "")

            doc.add_paragraph()
    elif not narrative:
        doc.add_paragraph(
            "（注：传导关系数据尚未生成。系统将在分析阶段自动生成环节间的"
            "成本传导、技术传导和价值传导路径。）"
        )


# ══════════════════════════════════════════════════════════════
# ── 第一章（下）：产业总体画像 ──
# ══════════════════════════════════════════════════════════════

def _add_chapter_profile(doc: Document, data: dict, narratives: dict):
    """第一章（续）：产业总体画像（市场数据），合并入第一章 全景概览"""
    # 合并入第一章：产业链全景概览（不单独设章标题）

    # AI 叙述（如有）
    narrative = narratives.get("industry_profile", "")
    if narrative:
        doc.add_paragraph(narrative)
        doc.add_paragraph()

    # 市场数据
    market_data = data.get("market_data", {})
    if market_data and any(v for v in market_data.values()):
        doc.add_heading("市场数据概览", level=2)

        if market_data.get("market_size"):
            p = doc.add_paragraph()
            run = p.add_run("市场规模：")
            run.font.bold = True
            p.add_run(market_data["market_size"])

        if market_data.get("forecast"):
            p = doc.add_paragraph()
            run = p.add_run("市场预测：")
            run.font.bold = True
            p.add_run(market_data["forecast"])

        if market_data.get("competition_landscape"):
            p = doc.add_paragraph()
            run = p.add_run("竞争格局：")
            run.font.bold = True
            p.add_run(market_data["competition_landscape"])

        if market_data.get("policy_environment"):
            p = doc.add_paragraph()
            run = p.add_run("政策环境：")
            run.font.bold = True
            p.add_run(market_data["policy_environment"])

        if market_data.get("key_drivers"):
            doc.add_paragraph("核心驱动因素：")
            for driver in market_data["key_drivers"]:
                doc.add_paragraph(driver, style="List Bullet")

        if market_data.get("opportunities"):
            doc.add_paragraph("发展机遇：")
            for opp in market_data["opportunities"]:
                doc.add_paragraph(opp, style="List Bullet")

        if market_data.get("risks"):
            doc.add_paragraph("风险因素：")
            for risk in market_data["risks"]:
                doc.add_paragraph(risk, style="List Bullet")


# ══════════════════════════════════════════════════════════════
# ── 第四章（上）：未来趋势推演 ──
# ══════════════════════════════════════════════════════════════

def _add_chapter_trends(doc: Document, data: dict, narratives: dict):
    """第四章：产业未来趋势与风险 = 原6 + 原8"""
    doc.add_heading("第四章  产业未来趋势与风险", level=1)

    doc.add_heading("未来趋势推演", level=2)

    narrative = narratives.get("trend_deduction", "")
    if narrative:
        _add_markdown_paragraphs(doc, narrative)
    else:
        doc.add_paragraph(
            '（注：趋势推演内容尚未生成。可通过前端的「AI 深度分析」功能，'
            '系统将基于产业链传导逻辑自动生成未来3-5年的趋势推演。）'
        )


# ── 第七章：衍生赛道补充分析（已按结构调整要求删除）──


def _add_chapter_risk(doc: Document, data: dict, narratives: dict):
    """第四章：风险与不确定性 = 原8，合并入第四章 未来趋势与风险"""
    doc.add_heading("风险与不确定性", level=2)

    narrative = narratives.get("risk_analysis", "")
    if narrative:
        _add_markdown_paragraphs(doc, narrative)
    else:
        doc.add_paragraph(
            '（注：风险与不确定性分析尚未生成。可通过前端的「AI 深度分析」功能，'
            '系统将基于产业链实际结构与传导逻辑，自动生成该产业特有的风险专题。）'
        )


# 证据状态 → 中文标签（与前端共用口径）
_EVENT_EVIDENCE_LABELS = {
    "VERIFIED": "已验证",
    "REPORTED": "已报道",
    "INFERRED": "推断",
    "UNVERIFIED": "未核实",
}


def _add_chapter_events(doc: Document, data: dict):
    """第五章：近期产业链动态 = 原9

    事件为「已验证/已报道的事实事件证据流」，不做利好/利空判断。
    每条事件列出日期、标签、摘要、涉及企业、证据状态与来源。
    """
    doc.add_heading("第五章  近期产业链动态", level=1)

    events = data.get("events", []) or []
    event_window = data.get("event_window", {}) or {}
    event_policy = data.get("event_policy", {}) or {}

    # 免责声明
    disclaimer = event_policy.get(
        "disclaimer",
        "本板块仅列已验证/已报道的事实事件，不构成任何投资建议；"
        "AI 不判断利好或利空，事件影响以原始披露为准。",
    )
    dp = doc.add_paragraph()
    run = dp.add_run(disclaimer)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True

    # 窗口元信息
    if event_window.get("window_days"):
        wp = doc.add_paragraph()
        wrun = wp.add_run(
            f"数据窗口：近 {event_window['window_days']} 天（基于公开研报）"
            + (f"｜生成时间：{event_window.get('generated_at', '')}" if event_window.get("generated_at") else "")
        )
        wrun.font.size = Pt(10)
        wrun.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if not events:
        doc.add_paragraph(
            "（注：近 "
            + str(event_window.get("window_days", "—"))
            + " 天（基于公开研报）未抽取到重大已验证产业链事件。"
            "本板块仅呈现有明确来源支撑的事实事件，不以空白填充。）"
        )
        return

    # 按日期倒序
    sorted_events = sorted(
        events,
        key=lambda e: (e.get("date") or ""),
        reverse=True,
    )

    for ev in sorted_events:
        if not isinstance(ev, dict):
            continue
        title = f"{ev.get('date', '日期未知')}　{ev.get('title', '（无标题）')}"
        doc.add_heading(title, level=2)

        # 标签 + 证据状态
        tags = ev.get("tags", []) or []
        ev_status = _EVENT_EVIDENCE_LABELS.get(ev.get("evidenceStatus", "REPORTED"), "已报道")
        meta_line = "｜".join([
            p for p in [
                ("标签：" + "、".join(tags)) if tags else "",
                f"证据状态：{ev_status}",
            ] if p
        ])
        if meta_line:
            mp = doc.add_paragraph()
            mrun = mp.add_run(meta_line)
            mrun.font.size = Pt(10)
            mrun.font.bold = True
            mrun.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

        # 摘要
        summary = ev.get("summaryZh", "")
        if summary:
            doc.add_paragraph(summary)

        # 涉及企业
        companies = ev.get("companies", []) or []
        if companies:
            doc.add_paragraph(f"涉及企业：{'、'.join(companies)}")

        # 影响护栏
        impact = ev.get("impactZh", "")
        if impact:
            ip = doc.add_paragraph()
            irun = ip.add_run(impact)
            irun.font.size = Pt(10)
            irun.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        # 来源
        source_title = ev.get("sourceTitle", "") or ev.get("sourceType", "研报")
        conf = ev.get("confidence")
        src_line = f"来源：{source_title}"
        if conf is not None:
            try:
                src_line += f"（置信度 {round(float(conf) * 100)}%）"
            except (ValueError, TypeError):
                pass
        sp = doc.add_paragraph()
        srun = sp.add_run(src_line)
        srun.font.size = Pt(10)
        srun.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()  # 段间距


# ══════════════════════════════════════════════════════════════
# ── 通用工具函数 ──
# ══════════════════════════════════════════════════════════════

# 行内加粗/强调标记：**bold** / __bold__ / *emphasis*
_INLINE_MD_RE = re.compile(r"\*\*(.+?)\*\*|__(.+?)__|\*(.+?)\*")


def _add_inline_runs(paragraph, text: str):
    """把含 **加粗** / *强调* 标记的文本渲染为带格式的 run，并去除残留星号。"""
    pos = 0
    for m in _INLINE_MD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        emph = m.group(1) or m.group(2) or m.group(3) or ""
        run = paragraph.add_run(emph)
        run.font.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])
    # 兜底：清除任何残留星号
    for run in paragraph.runs:
        if "*" in run.text:
            run.text = run.text.replace("*", "")


def _add_markdown_paragraphs(doc: Document, text: str):
    """将 Markdown 文本渲染为 Word 段落：

    - ``**加粗**`` / ``*强调*`` → Word 加粗
    - 以 ``*`` / ``-`` / ``+`` 开头的行 → Word 项目符号
    - ``#`` 标题行 → 加粗小标题
    - 去除任何残留星号
    """
    if not text:
        return
    for raw in text.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        # 标题行 #～######
        mh = re.match(r"^\s*#{1,6}\s+(.*)$", line)
        if mh:
            p = doc.add_paragraph()
            run = p.add_run(mh.group(1).strip().replace("*", ""))
            run.font.bold = True
            run.font.size = Pt(11)
            continue
        # 项目符号行 * / - / +
        mb = re.match(r"^\s*[\*\-\+]\s+(.*)$", line)
        if mb:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, mb.group(1).strip())
            continue
        # 普通段落
        p = doc.add_paragraph()
        _add_inline_runs(p, line.strip())


def _add_chart(doc: Document, image_path: str, caption: str, width=Inches(5.5)):
    """在文档中嵌入图表并添加标题"""
    if not os.path.exists(image_path):
        logger.warning(f"图表文件不存在: {image_path}")
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=width)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True


def _add_disclaimer(doc: Document):
    """免责声明"""
    doc.add_paragraph()
    doc.add_paragraph()
    disclaimer = doc.add_paragraph(
        "免责声明：本报告由 AI 系统基于公开研报自动生成，仅供产业研究参考。"
        "报告中引用的数据和观点版权归原作者及券商所有。"
        "本报告不构成任何投资建议。"
    )
    disclaimer.runs[0].font.size = Pt(9)
    disclaimer.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _add_key_companies(doc: Document, data: dict):
    """附录 · 重点企业分析（与前端 / HTML 附录一致：汇总产业链重点上市企业）"""
    companies = data.get("companies") or []

    def ok(c):
        if not (c.get("name") or "").strip():
            return False
        if not (c.get("stock_code") or "").strip():
            return False
        if not (c.get("chain_position") or "").strip():
            return False
        if not (c.get("sub_segment") or "").strip():
            return False
        if not (c.get("main_business") or "").strip():
            return False
        prods = c.get("products") or []
        if not isinstance(prods, list) or len(prods) == 0:
            return False
        if all((not (p or "").strip()) for p in prods):
            return False
        return True

    filtered = [c for c in companies if ok(c)]
    if not filtered:
        return

    doc.add_heading("附录 · 重点企业分析", level=1)
    doc.add_paragraph("本节汇总产业链中的重点上市企业及其定位、主营业务与核心产品。")
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(
        ["企业名称", "股票代码", "产业链位置", "细分环节", "主营业务", "核心产品"]
    ):
        hdr[i].text = h
    for c in filtered[:30]:
        prods = c.get("products") or []
        prod_str = "、".join([p for p in prods[:3] if (p or "").strip()])
        row = table.add_row().cells
        row[0].text = c.get("name") or ""
        row[1].text = c.get("stock_code") or ""
        row[2].text = c.get("chain_position") or ""
        row[3].text = c.get("sub_segment") or ""
        row[4].text = (c.get("main_business") or "")[:80]
        row[5].text = prod_str
    if len(filtered) > 30:
        doc.add_paragraph(f"共 {len(filtered)} 家企业，仅展示前 30 家。")


def _add_relations(doc: Document, data: dict):
    """附录 · 企业关系（与前端 / HTML 附录一致：仅保留能匹配到公开研报佐证的关系）"""
    relations = data.get("relations") or []
    reference_reports = data.get("_reference_reports") or []

    def has_related(r):
        targets = [x for x in (r.get("from_company"), r.get("to_company")) if x]
        for rep in reference_reports:
            hay = f"{rep.get('title') or ''} {rep.get('stock_name') or ''}"
            if any(t and t in hay for t in targets):
                return True
        return False

    visible = [r for r in relations if has_related(r)]
    if not visible:
        return

    doc.add_heading("附录 · 企业关系", level=1)
    doc.add_paragraph("本节列出产业链企业间的供需与关联关系（仅展示能匹配到公开研报佐证的关系）。")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["供应方", "需求方", "关系类型", "详情"]):
        hdr[i].text = h
    for r in visible[:50]:
        row = table.add_row().cells
        row[0].text = r.get("from_company") or "-"
        row[1].text = r.get("to_company") or "-"
        row[2].text = r.get("type") or "-"
        row[3].text = (r.get("detail") or "")[:80]
    if len(visible) > 50:
        doc.add_paragraph(f"共 {len(visible)} 条关系，仅展示前 50 条。")
